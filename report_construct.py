from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_UNDERLINE
from docx.shared import Cm, Pt, Inches
from docx.enum.section import WD_ORIENT
from docx import Document

from datetime import datetime
from babel.dates import format_date, format_datetime, format_time
from babel import Locale


def string_wrapper(input_string: str, max_length: int = 30, symbol: str = "_") -> str:
    """Surround string with chosen symbols to fit max_length

    :param input_string: string to wrap
    :type input_string: str
    :param max_length: maximum string length to wrap, defaults to 35
    :type max_length: int, optional
    :param symbol: symbol to wrap, defaults to "_"
    :type symbol: str, optional
    :return: wrapped string
    :rtype: str

    For example string_wrapper("name", 10, "_") -> ___name___
    """
    # check if string is empty
    if not input_string:
        return ""
    # check if string is bigger then max_lenght
    surrounding_length = max_length - len(input_string)

    if surrounding_length <= 0:
        return input_string

    # calc amount of symbols for each side
    left_length = surrounding_length // 2
    right_length = surrounding_length - left_length

    # Form new string itself
    surrounded_string = symbol * left_length + input_string + symbol * right_length

    return surrounded_string


def split_string(long_string: str, max_length: int = 35) -> list[str]:
    """Func to split long str into smaller with lenght <= max_lenght

    :param long_string: initial string
    :type long_string: str
    :param max_length: maximum symbols per string, defaults to 35
    :type max_length: int
    :return: list of strings after splitting long_string
    :rtype: list[str]
    """
    # check if string is exist
    if not long_string:
        return []
    # list to store res
    short_strings = []

    # list of words in long_string
    words = long_string.split()

    current_string = words[0]
    for word in words[1:]:
        # if possible, we extend current sting with another word
        if len(current_string) + len(word) + 1 <= max_length:
            current_string += " " + word
        else:
            # if string cannot be bigger, we add it to result and annul current string
            short_strings.append(current_string)
            current_string = word
    # last string goes in
    short_strings.append(current_string)

    return short_strings


def date_formatter(date_string: str):
    """Date formatter

    :param date_string: date in any day.month.year
    :type date_string: str
    :return: date string in format day month(word) year
    :rtype: str
    """
    locale = Locale("ru", "RU")
    # converting string to datetime
    date_object = datetime.strptime(date_string, "%d.%m.%Y")

    # making good looking string as i need
    formatted_date = format_date(date_object, format="«d» MMMM y г.", locale=locale)

    return formatted_date


def add_header(doc: Document, con_position: str, name: str, date: str):
    """upper right corner info about report

    :param doc: doument itself
    :type doc: Document
    :param con_position: responsible person position
    :type con_position: str
    :param name: responsible person name
    :type name: str
    :param date: sign date
    :type date: str
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run("Утверждаю")
    position_strs = split_string(con_position)
    for p_str in position_strs:
        paragraph = doc.add_paragraph()
        paragraph.alignment = 2
        run = paragraph.add_run(p_str)
    wrapped_name = string_wrapper(name)

    paragraph = doc.add_paragraph()
    paragraph.alignment = 2
    run = paragraph.add_run(wrapped_name)
    run.font.underline = True

    paragraph = doc.add_paragraph()
    paragraph.alignment = 2
    run = paragraph.add_run(date_formatter(date))
    run.font.underline = True

    paragraph = doc.add_paragraph()
    paragraph.alignment = 2
    run = paragraph.add_run("\t М.П.")

    hollow_line(doc)


def add_signatures(doc: Document, members: list[list[str]]) -> None:
    """signature fields like
    position      name      (field for signature)

    :param doc: document
    :type doc: Document
    :param members: list of all signature fields
    :type members: list[list[str]]
    """
    for member in members:
        paragraph = doc.add_paragraph()
        paragraph.alignment = 0  # WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run(
            "_".join(
                [string_wrapper(member[0], 40), string_wrapper(member[1], 40), "_" * 40]
            )
        )
        run.font.underline = True

        paragraph = doc.add_paragraph()
        paragraph.alignment = 0
        run = paragraph.add_run(
            " ".join(
                string_wrapper(caption, 85, " ")
                for caption in ["(должность)", "(ФИО)", "(подпись)"]
            )
        )
        run.font.size = Pt(10)

        hollow_line(doc)


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_table(
    doc: Document,
    data: list[list[str]],
    col_widths: list[Cm] = [
        Cm(1.3),
        Cm(2.5),
        Cm(2.7),
        Cm(5.7),
        Cm(1.7),
        Cm(1.7),
        Cm(11.2),
    ],
) -> None:
    """add table with given data into doc

    :param doc: document itself
    :type doc: Document
    :param data: list with list for each row
    :type data: list[list[str]]
    :param col_widths: Optional, setting for column widths, defaults to [ Inches(0.6), Inches(1.2), Inches(1.2), Inches(2.2), Inches(0.7), Inches(0.7), Inches(4), ]
    :type col_widths: list[Cm], optional
    """
    # Add the table
    num_rows = len(data)
    num_cols = len(data[0])
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"
    # Set column widths
    set_col_widths(table, col_widths)
    # table.autofit = False
    # for i, col_width in enumerate(col_widths):
    #     table.columns[i].width = col_width
    for i in range(num_rows):
        for j in range(num_cols):
            cell = table.cell(i, j)
            cell.text = data[i][j]
    # Set font size for table cells
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    hollow_line(doc)


def hollow_line(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = 0
    run = paragraph.add_run("")


def risk_map_report_constructor(
    prefix_data: list[str],
    map_num: int,
    enterprise: str,
    subdistrict: list[str],
    process: str,
    table_data: list[list[str]],
    map_creator_list: list[list[str]],
    commition_list: list[list[str]],
    acquatence_list: list[list[str]],
    file_path: str,
):
    # Create a new document
    doc = Document()
    ### doc settings
    # Set page size to A4
    section = doc.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    # Set document orientation to landscape

    # Set margins
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # Font and text size
    font = doc.styles["Normal"].font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Line spacing
    paragraph_format = doc.styles["Normal"].paragraph_format
    paragraph_format.line_spacing = 1

    # Indentation and spacing
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = 0

    # filling header
    add_header(doc, *prefix_data)

    # map number
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(f"Карта оценки рисков № {map_num}")

    hollow_line(doc)
    # filling pre table annotations
    paragraph = doc.add_paragraph()
    paragraph.alignment = 0  # WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run("Организация: ")
    run = paragraph.add_run(string_wrapper(enterprise, 105))
    run.font.underline = True

    hollow_line(doc)

    paragraph = doc.add_paragraph()
    paragraph.alignment = 0
    run = paragraph.add_run("Подразделение: ")
    run = paragraph.add_run(string_wrapper(", ".join(subdistrict), 105))
    run.font.underline = True

    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run("(должность, место работы)")
    run.font.size = Pt(10)

    hollow_line(doc)

    paragraph = doc.add_paragraph()
    paragraph.alignment = 0
    run = paragraph.add_run("Процессы: ")
    run = paragraph.add_run(string_wrapper(process, 109))
    run.font.underline = True
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run("(описание работ)")
    run.font.size = Pt(10)
    hollow_line(doc)
    add_table(doc, table_data)
    # adding signs

    paragraph = doc.add_paragraph()
    paragraph.alignment = 0
    run = paragraph.add_run("Карту составил:")
    add_signatures(doc, map_creator_list)

    hollow_line(doc)

    paragraph = doc.add_paragraph()
    paragraph.alignment = 0
    run = paragraph.add_run("Члены коммисии:")
    add_signatures(doc, commition_list)

    hollow_line(doc)

    paragraph = doc.add_paragraph()
    paragraph.alignment = 0
    run = paragraph.add_run("С картой ознакомлен:")
    add_signatures(doc, acquatence_list)

    doc.save(file_path)


# def fill_document(file_path, data):
#     # Create a new document
#     doc = Document()

#     # Set page size to A4
#     section = doc.sections[0]

#     section.orientation = WD_ORIENT.LANDSCAPE
#     section.page_width = Cm(29.7)
#     section.page_height = Cm(21)
#     # Set document orientation to landscape

#     # Set margins
#     section.left_margin = Cm(1.5)
#     section.right_margin = Cm(1.5)
#     section.top_margin = Cm(1)
#     section.bottom_margin = Cm(1)

#     # Font and text size
#     font = doc.styles["Normal"].font
#     font.name = "Times New Roman"
#     font.size = Pt(12)

#     # Line spacing
#     paragraph_format = doc.styles["Normal"].paragraph_format
#     paragraph_format.line_spacing = 1

#     # Indentation and spacing
#     paragraph_format.space_before = Pt(0)
#     paragraph_format.space_after = Pt(0)
#     paragraph_format.line_spacing_rule = 1

#     add_header(doc, data[1], data[2], data[3])
#     # # Add the first 4 right-aligned lines
#     # for i in range(4):
#     #     paragraph = doc.add_paragraph()
#     #     paragraph.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.RIGHT
#     #     run = paragraph.add_run(data[i])

#     # Add the centered line
#     paragraph = doc.add_paragraph()
#     paragraph.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
#     run = paragraph.add_run(data[4])
#     run.font.size = Pt(10)
#     # Add the table
#     data = data[5]
#     num_rows = len(data)
#     num_cols = len(data[0])
#     table = doc.add_table(rows=num_rows, cols=num_cols)
#     table.style = "Table Grid"

#     # Set column widths

#     col_widths = [
#         Inches(0.6),
#         Inches(1.2),
#         Inches(1.2),
#         Inches(2.2),
#         Inches(0.7),
#         Inches(0.7),
#         Inches(4),
#     ]
#     for i, col_width in enumerate(col_widths):
#         table.columns[i].width = col_width
#     table.autofit = True
#     # Fill the table with data
#     for i in range(num_rows):
#         for j in range(num_cols):
#             cell = table.cell(i, j)
#             cell.text = data[i][j]
#     # Set font size for table cells
#     for row in table.rows:
#         for cell in row.cells:
#             cell.paragraphs[0].paragraph_format.space_before = Pt(0)
#             cell.paragraphs[0].paragraph_format.space_after = Pt(0)
#             for paragraph in cell.paragraphs:
#                 for run in paragraph.runs:
#                     run.font.size = Pt(10)
#     # Add the constructions
#     for i in range(6, len(data) - 1):
#         paragraph = doc.add_paragraph(data[i])
#         paragraph.alignment = 0  # WD_PARAGRAPH_ALIGNMENT.LEFT
#     members = [
#         ["Командир отделения холодильных камер", "Холодный И.В."],
#         ["Начальник по управлению туалетом", "Приморский Т.П."],
#     ]
#     # Add the signature construction with underline
#     add_signatures(doc, members)
#     # paragraph = doc.add_paragraph(data[-1])
#     # paragraph.alignment = 0  # WD_PARAGRAPH_ALIGNMENT.LEFT
#     # run = paragraph.runs[0]
#     # run.font.underline = True

#     # Save the document
#     doc.save(file_path)

if __name__ == "__main__":
    data = [
        "Утверждаю",
        "Директор <<оао шляпи>>",
        "Ковалёв Р.В.",
        "15.1.2022",
        "Карта оценки рисков № _",
        [
            [
                "Код риска",
                "Источник опасности",
                "Опасность",
                "Возможные последствия",
                "Вероятность риска",
                "Ущерб",
                "Меры по снижению риска",
            ],
            [
                "1",
                "Резервуар для хранения СПГ, трубопроводы ",
                "Утечка газообразного азота",
                "Утечка газообразного азота может привести к образованию атмосферы, снижающей содержание кислорода, что создает риск асфиксии для персонала ",
                "Средняя ",
                "Средний ",
                "a. Регулярное техническое обслуживание и проверка герметичности азотных систем и соединений;\n b. Обучение персонала правилам эксплуатации и обслуживания азотных систем, включая меры безопасности и процедуры аварийного реагирования \nc. Установка датчиков кислорода и сигнализации для контроля содержания кислорода в рабочих зонах; \nd. Использование специальных средств индивидуальной защиты для работы с азотом, включая кислородные маски. ",
            ],
            [
                "2",
                "Источник 2",
                "Опасность 2",
                "Последствия 2",
                "0.5",
                "Ущерб 2",
                "Меры 2",
            ],
        ],
        "Сделал это:",
        "Начальник отдела        Шляпик В.И.         ___________",
        "(должность)    (ФИО)     (подпись)",
    ]
    prefix_data = [
        "Директор ОАО МежГазРегионСтрой южного округа",
        "Персиков В.В.",
        "12.01.2022",
    ]
    map_num = 1
    enterprise = "ОАО МежГазРегионСтрой южного округа"
    subdistrict = ["Интернет-герой", "Рота диванных войск №5"]
    process = "Интернет атаки на инагентов США"
    table_data = [
        [
            "Код риска",
            "Источник опасности",
            "Опасность",
            "Возможные последствия",
            "Вероятность риска",
            "Ущерб",
            "Меры по снижению риска",
        ],
        [
            "1",
            "Резервуар для хранения СПГ, трубопроводы ",
            "Утечка газообразного азота",
            "Утечка газообразного азота может привести к образованию атмосферы, снижающей содержание кислорода, что создает риск асфиксии для персонала ",
            "Средняя ",
            "Средний ",
            "a. Регулярное техническое обслуживание и проверка герметичности азотных систем и соединений;\n b. Обучение персонала правилам эксплуатации и обслуживания азотных систем, включая меры безопасности и процедуры аварийного реагирования \nc. Установка датчиков кислорода и сигнализации для контроля содержания кислорода в рабочих зонах; \nd. Использование специальных средств индивидуальной защиты для работы с азотом, включая кислородные маски. ",
        ],
        [
            "2",
            "Источник 2",
            "Опасность 2",
            "Последствия 2",
            "0.5",
            "Ущерб 2",
            "Меры 2",
        ],
    ]
    map_creator_list = [["Начальник отдела по расследованиям", "Бобик Ч.М."]]
    commition_list = [
        ["Царь руси по линии наследования", "Грозный И.В."],
        ["Министр внутренней ереси", "Шило В.П."],
    ]
    acquatence_list = [["Директор интернета", "Комаров И.И."]]
    file_path = "заполненный_документ.docx"
    risk_map_report_constructor(
        prefix_data,
        map_num,
        enterprise,
        subdistrict,
        process,
        table_data,
        map_creator_list,
        commition_list,
        acquatence_list,
        file_path,
    )

    # fill_document("заполненный_документ.docx", data)
